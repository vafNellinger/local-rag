"""Dokument → Text, formatabhängig.

Zwei Stufen, weil sie zehnfach unterschiedlich teuer sind:

``probe()`` liest den Text-Layer mit pypdf und stellt fest, welche Seiten
überhaupt OCR brauchen. Kostet ~0,05 s pro Datei.

``convert()`` macht die eigentliche Extraktion über Docling: Layout-Analyse,
Tabellen, Überschriften, bei Bedarf OCR. Gemessen ~0,8 s pro Seite ohne OCR
und ~8,5 s pro Seite mit OCR — deshalb entscheidet ``probe()`` vorab, ob OCR
eingeschaltet wird, statt es pauschal laufen zu lassen.

Der interessante Teil sind PDFs. Ein PDF ist kein Format, sondern zwei:
digital erzeugte Seiten mit Text-Layer und gescannte Seiten, die nur ein Bild
enthalten. Beides kommt in derselben Datei vor — gescanntes Deckblatt vor
digitalem Rest ist der Normalfall. Die Entscheidung fällt deshalb pro Seite.
"""

from __future__ import annotations

import hashlib
import json
import logging
import time
import unicodedata
import warnings
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

# Sprachen für die Texterkennung. Deutsch zuerst, Englisch dazu, weil
# Fachbegriffe und Produktnamen in deutschen Dokumenten regelmäßig englisch
# sind. Am erzeugten Testscan verifiziert: Umlaute und ß werden korrekt
# erkannt ("Änderungsantrag", "Größen", "gemäß", "Jürgen Öztürk").
DEFAULT_OCR_LANGS = ("de", "en")

# Formate, die Docling verarbeitet. Für Markdown und Text lohnt der Aufwand
# nicht — die sind bereits Text und werden direkt gelesen.
DOCLING_SUFFIXES = {".pdf", ".docx"}

# Ab wie vielen extrahierten Zeichen eine PDF-Seite als "hat Text-Layer" gilt.
# Gescannte Seiten liefern typischerweise 0 Zeichen, manchmal ein paar
# Artefakte aus eingebetteten Wasserzeichen oder Seitenzahlen. 100 trennt das
# sauber, ohne echte, aber dünn belegte Seiten (Titelblatt, Kapiteltrenner)
# fälschlich als Scan zu werten.
TEXT_LAYER_MIN_CHARS = 100

# Zweites Kriterium: eine Seite mit viel Text, der zu großen Teilen aus
# Ersatzzeichen besteht, hat eine kaputte Zeichenkodierung (fehlendes
# ToUnicode-CMap). Das sieht wie Text aus, ist aber unbrauchbar — solche
# Seiten gehören ebenfalls ins OCR.
MAX_REPLACEMENT_RATIO = 0.15

# Drittes Kriterium, und das entscheidende: OCR lohnt nur, wenn es ein Bild
# gibt. Eine Seite mit wenig Text und ohne Bild ist dünn belegt (Schlussformel,
# Kapiteltrenner), nicht gescannt — OCR würde dort ein leeres Blatt abtasten.
#
# Die Mindestfläche trennt Seitenscans von Logos und Signaturbildchen. Ein
# Scan deckt die Seite ab und liegt bei mehreren Millionen Pixeln; ein Logo
# in der Kopfzeile bleibt weit darunter.
MIN_SCAN_IMAGE_PIXELS = 100_000

# Wie tief in verschachtelte Form-XObjects hinein nach Bildern gesucht wird.
# Zwei Ebenen decken die üblichen Verschachtelungen ab, ohne bei bösartig
# verketteten Dokumenten in eine Endlosschleife zu laufen.
MAX_XOBJECT_DEPTH = 2

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".markdown", ".txt", ".text"}

# Wohin extrahiertes Markdown gelegt wird, damit ein zweiter Durchgang die
# teure Stufe überspringen kann.
#
# Der Cache liegt ausdrücklich *außerhalb* des Index. Der Anlass ist ein
# Wechsel des Embedding-Modells: dabei wird der Index verworfen und neu
# gebaut, das Extraktionsergebnis bleibt aber identisch. Läge der Text im
# Index, wäre er genau dann unerreichbar, wenn man ihn braucht — ein Index mit
# fremdem Modell lässt sich nicht öffnen.
#
# Gemessen am Testkorpus macht die Extraktion ein Drittel des Ingests aus, bei
# 10 % gescannten Seiten über die Hälfte.
EXTRACT_CACHE_DIR = Path.home() / ".cache" / "local-rag" / "extract"

# Hochzählen, wenn sich das Format der Cache-Einträge ändert. Alte Einträge
# werden dann ignoriert statt falsch gelesen.
EXTRACT_CACHE_VERSION = 1

# Formate, die eine Konvertierung brauchen, bevor wir sie lesen können.
# Explizit benannt, damit die Fehlermeldung den Weg nennt statt nur "geht nicht".
CONVERSION_HINTS = {
    ".doc": "altes Word-Format — mit libreoffice --convert-to docx umwandeln",
    ".ppt": "altes PowerPoint-Format — mit libreoffice --convert-to pptx umwandeln",
    ".xls": "altes Excel-Format — mit libreoffice --convert-to xlsx umwandeln",
    ".pages": "Apple Pages — als PDF oder DOCX exportieren",
}


class ExtractionError(RuntimeError):
    """Die Datei konnte nicht gelesen werden."""


@dataclass
class Page:
    """Eine Seite oder ein Abschnitt eines Dokuments.

    ``status`` unterscheidet die Fälle, die sonst alle als "kein Text"
    zusammenfielen:

    - ``text``   — Text-Layer ist brauchbar
    - ``scan``   — kein brauchbarer Text, aber ein seitenfüllendes Bild: OCR hilft
    - ``sparse`` — kein Text und kein Bild: die Seite ist wirklich leer, OCR
      würde nichts finden
    - ``error``  — Seite nicht lesbar
    """

    number: int
    text: str
    status: str = "text"
    largest_image_pixels: int = 0
    # Gefüllt, sobald ein OCR-Backend die Seite bearbeitet hat.
    ocr_applied: bool = False
    note: str | None = None

    @property
    def needs_ocr(self) -> bool:
        return self.status == "scan" and not self.ocr_applied

    @property
    def char_count(self) -> int:
        return len(self.text.strip())


@dataclass
class Document:
    """Ergebnis der Extraktion."""

    path: Path
    format: str
    pages: list[Page] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def char_count(self) -> int:
        return sum(p.char_count for p in self.pages)

    @property
    def pages_needing_ocr(self) -> list[Page]:
        return [p for p in self.pages if p.needs_ocr]

    @property
    def sparse_pages(self) -> list[Page]:
        """Seiten ohne Text und ohne Bild — dort ist nichts zu holen."""
        return [p for p in self.pages if p.status == "sparse"]

    @property
    def ocr_ratio(self) -> float:
        """Anteil der Seiten, die ohne OCR nicht lesbar sind."""
        if not self.pages:
            return 0.0
        return len([p for p in self.pages if p.status == "scan"]) / len(self.pages)


def _replacement_ratio(text: str) -> float:
    """Anteil an Ersatz- und Steuerzeichen — Indikator für kaputte Kodierung."""
    stripped = text.strip()
    if not stripped:
        return 0.0
    bad = sum(
        1
        for ch in stripped
        if ch == "�" or (unicodedata.category(ch) == "Cc" and ch not in "\n\r\t")
    )
    return bad / len(stripped)


def _largest_image_pixels(page_obj, depth: int = 0) -> int:
    """Größtes Bild auf einer PDF-Seite, in Pixeln. 0 wenn keins.

    Sucht auch in Form-XObjects, weil Scanner-Software das Seitenbild gern in
    einen Form-Container verpackt statt direkt einzubetten.
    """
    if depth > MAX_XOBJECT_DEPTH:
        return 0

    try:
        resources = page_obj.get("/Resources")
        if resources is None:
            return 0
        xobjects = resources.get_object().get("/XObject")
        if xobjects is None:
            return 0
        xobjects = xobjects.get_object()
    except Exception as exc:
        logger.debug("XObject-Zugriff fehlgeschlagen: %s", exc)
        return 0

    largest = 0
    for key in xobjects:
        try:
            obj = xobjects[key].get_object()
            subtype = obj.get("/Subtype")
            if subtype == "/Image":
                pixels = int(obj.get("/Width") or 0) * int(obj.get("/Height") or 0)
                largest = max(largest, pixels)
            elif subtype == "/Form":
                largest = max(largest, _largest_image_pixels(obj, depth + 1))
        except Exception as exc:
            logger.debug("Bild %s nicht auswertbar: %s", key, exc)
            continue

    return largest


def _classify_page(text: str, image_pixels: int) -> tuple[str, str | None]:
    """Bestimme den Seitenstatus. Gibt (status, Grund).

    Die Bildprüfung ist ausschlaggebend: ohne seitenfüllendes Bild gibt es
    nichts zu erkennen, egal wie leer der Text-Layer ist.
    """
    stripped = text.strip()
    has_scan_image = image_pixels >= MIN_SCAN_IMAGE_PIXELS

    if len(stripped) < TEXT_LAYER_MIN_CHARS:
        if has_scan_image:
            return "scan", f"{len(stripped)} Zeichen, Bild mit {image_pixels:,} px"
        return "sparse", f"nur {len(stripped)} Zeichen, kein Bild — Seite ist leer"

    ratio = _replacement_ratio(stripped)
    if ratio > MAX_REPLACEMENT_RATIO:
        if has_scan_image:
            return "scan", f"{ratio:.0%} Ersatzzeichen, Bild vorhanden"
        # Kaputte Kodierung ohne Bild: OCR kann nicht helfen, aber der Text
        # ist unbrauchbar. Muss sichtbar bleiben statt still durchzurutschen.
        return "error", f"{ratio:.0%} Ersatzzeichen, kein Bild — Kodierung defekt"

    return "text", None


def _extract_pdf(path: Path) -> Document:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("pypdf fehlt: uv pip install pypdf") from exc

    doc = Document(path=path, format="pdf")

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ExtractionError(f"PDF nicht lesbar: {exc}") from exc

    if reader.is_encrypted:
        # Viele PDFs sind mit leerem Passwort verschlüsselt — das lässt sich
        # ohne Zutun öffnen. Erst wenn auch das scheitert, ist wirklich Schluss.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError(f"PDF ist passwortgeschützt: {exc}") from exc
        doc.warnings.append("PDF war verschlüsselt, mit leerem Passwort geöffnet")

    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            # Eine kaputte Seite darf nicht das ganze Dokument verlieren.
            # Als "scan" markiert, weil OCR hier die einzige Chance ist.
            logger.debug("Seite %d nicht extrahierbar: %s", index, exc)
            doc.pages.append(
                Page(
                    number=index,
                    text="",
                    status="scan",
                    note=f"Text-Layer defekt ({exc}) — OCR als Rückfallebene",
                )
            )
            continue

        image_pixels = _largest_image_pixels(page)
        status, reason = _classify_page(text, image_pixels)
        doc.pages.append(
            Page(
                number=index,
                text=text,
                status=status,
                largest_image_pixels=image_pixels,
                note=reason,
            )
        )

    return doc


def _extract_docx(path: Path) -> Document:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("python-docx fehlt: uv pip install python-docx") from exc

    try:
        source = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError(f"DOCX nicht lesbar: {exc}") from exc

    doc = Document(path=path, format="docx")

    blocks = [p.text for p in source.paragraphs if p.text.strip()]

    # Tabellen stehen in DOCX außerhalb des Absatzflusses und gingen sonst
    # verloren. Zeilenweise mit Tabs zusammengesetzt bleibt die Zuordnung
    # Spalte-zu-Wert für das Sprachmodell erkennbar.
    table_count = 0
    for table in source.tables:
        table_count += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))

    if table_count:
        doc.warnings.append(
            f"{table_count} Tabelle(n) zeilenweise eingefügt — "
            "Layout geht verloren, Inhalt bleibt"
        )

    # DOCX hat keine verlässlichen Seitengrenzen (die entstehen erst beim
    # Rendern), deshalb ein Block.
    doc.pages.append(Page(number=1, text="\n".join(blocks)))
    return doc


def _extract_text(path: Path, fmt: str) -> Document:
    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback für Dateien aus Windows-Umgebungen.
        content = path.read_text(encoding="latin-1")
        doc = Document(path=path, format=fmt)
        doc.warnings.append("Datei war nicht UTF-8, als latin-1 gelesen")
        doc.pages.append(Page(number=1, text=content))
        return doc

    doc = Document(path=path, format=fmt)
    doc.pages.append(Page(number=1, text=content))
    return doc


def probe(path: str | Path) -> Document:
    """Schnelle Analyse: was steckt drin und welche Seiten brauchen OCR.

    Führt kein OCR aus und macht keine Layout-Analyse — das ist Aufgabe von
    ``convert()``. Diese Stufe existiert, um die teure Stufe gezielt zu
    steuern.
    """
    file_path = Path(path).expanduser()

    if not file_path.exists():
        raise ExtractionError(f"Datei nicht gefunden: {file_path}")
    if not file_path.is_file():
        raise ExtractionError(f"Kein Dateipfad: {file_path}")

    suffix = file_path.suffix.lower()

    if hint := CONVERSION_HINTS.get(suffix):
        raise ExtractionError(f"{suffix} wird nicht direkt unterstützt — {hint}")

    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix in {".md", ".markdown"}:
        return _extract_text(file_path, "markdown")
    if suffix in {".txt", ".text"}:
        return _extract_text(file_path, "text")

    raise ExtractionError(
        f"Format {suffix or '(ohne Endung)'} wird nicht unterstützt. "
        f"Möglich: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
    )


# ─── Stufe 2: Docling ────────────────────────────────────────────────────────

# DocumentConverter lädt beim ersten Gebrauch Layout- und OCR-Modelle. Das
# kostet zweistellige Sekunden, deshalb pro Konfiguration genau eine Instanz.
_CONVERTER_CACHE: dict[tuple, object] = {}


@dataclass
class ConvertedDocument:
    """Ergebnis der vollständigen Extraktion."""

    path: Path
    format: str
    markdown: str
    ocr_used: bool
    page_count: int
    duration_seconds: float
    warnings: list[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.markdown)


def _get_converter(*, ocr: bool, ocr_langs: tuple[str, ...]):
    key = (ocr, ocr_langs)
    if key in _CONVERTER_CACHE:
        return _CONVERTER_CACHE[key]

    # Torch meldet bei jedem DataLoader, dass ohne Beschleuniger kein
    # pin_memory genutzt wird. Auf einer CPU-Installation ist das der
    # Normalzustand und keine Information, die den Anwender erreichen muss.
    warnings.filterwarnings("ignore", message=".*pin_memory.*")

    try:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import (
            EasyOcrOptions,
            PdfPipelineOptions,
        )
        from docling.document_converter import DocumentConverter, PdfFormatOption
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError(
            "docling fehlt: uv pip install docling easyocr"
        ) from exc

    options = PdfPipelineOptions()
    options.do_ocr = ocr
    options.do_table_structure = True
    if ocr:
        # Explizit EasyOCR statt der Auto-Wahl: die Standardeinstellung von
        # Docling lässt die Sprachliste leer, und ohne 'de' leiden Umlaute.
        options.ocr_options = EasyOcrOptions(lang=list(ocr_langs))

    converter = DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=options)}
    )
    _CONVERTER_CACHE[key] = converter
    return converter


def _file_hash(path: Path) -> str:
    """Hash über den Dateiinhalt, als Cache-Schlüssel.

    Eigene Kopie und nicht ``store.file_sha256``: die Extraktion soll den
    Index nicht kennen. Inhalt statt mtime, weil Kopieren und Auspacken die
    mtime neu setzen, ohne dass sich etwas geändert hat.
    """
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        while block := handle.read(1024 * 1024):
            digest.update(block)
    return digest.hexdigest()


def _docling_version() -> str:
    """Version der Extraktionsbibliothek, als Teil des Cache-Schlüssels.

    Eine neue Docling-Version liefert anderes Layout — ein Cache, der dann
    alte Ergebnisse zurückgibt, wäre ein stiller Fehler. Nach einem Update
    wird also neu extrahiert.
    """
    try:
        from importlib.metadata import version

        return version("docling")
    except Exception:  # pragma: no cover
        return "unbekannt"


def _cache_key(file_hash: str, *, ocr: bool, ocr_langs: tuple[str, ...]) -> str:
    """Schlüssel eines Cache-Eintrags.

    Der OCR-Modus gehört hinein: dieselbe Datei ergibt mit und ohne
    Texterkennung ein anderes Markdown, und die Sprachliste beeinflusst das
    Ergebnis ebenfalls.
    """
    parts = [
        str(EXTRACT_CACHE_VERSION),
        file_hash,
        "ocr" if ocr else "kein-ocr",
        "-".join(ocr_langs),
        _docling_version(),
    ]
    digest = hashlib.sha256("|".join(parts).encode()).hexdigest()
    return digest[:32]


def _cache_path(key: str) -> Path:
    return EXTRACT_CACHE_DIR / f"{key}.json"


def _read_cache(key: str) -> dict | None:
    path = _cache_path(key)
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as exc:
        logger.debug("Cache-Eintrag %s unlesbar: %s", path.name, exc)
        return None


def _write_cache(key: str, payload: dict) -> None:
    try:
        EXTRACT_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        _cache_path(key).write_text(
            json.dumps(payload, ensure_ascii=False), encoding="utf-8"
        )
    except OSError as exc:
        # Ein nicht schreibbarer Cache darf die Extraktion nicht scheitern
        # lassen — er ist eine Abkürzung, keine Voraussetzung.
        logger.warning("Cache konnte nicht geschrieben werden: %s", exc)


def clear_extract_cache() -> int:
    """Alle Cache-Einträge löschen. Gibt die Zahl der entfernten Dateien."""
    if not EXTRACT_CACHE_DIR.exists():
        return 0
    entfernt = 0
    for entry in EXTRACT_CACHE_DIR.glob("*.json"):
        try:
            entry.unlink()
            entfernt += 1
        except OSError as exc:  # pragma: no cover
            logger.debug("%s nicht löschbar: %s", entry, exc)
    return entfernt


def extract_cache_size() -> tuple[int, int]:
    """(Zahl der Einträge, Bytes) des Extraktions-Caches."""
    if not EXTRACT_CACHE_DIR.exists():
        return 0, 0
    entries = list(EXTRACT_CACHE_DIR.glob("*.json"))
    return len(entries), sum(e.stat().st_size for e in entries)


def convert(
    path: str | Path,
    *,
    ocr: bool | None = None,
    ocr_langs: tuple[str, ...] = DEFAULT_OCR_LANGS,
    use_cache: bool = True,
    file_hash: str | None = None,
) -> ConvertedDocument:
    """Extrahiere ein Dokument als strukturiertes Markdown.

    ``ocr=None`` entscheidet automatisch über ``probe()``: OCR läuft nur, wenn
    mindestens eine Seite als Scan erkannt wurde. Das ist der Unterschied
    zwischen 29 Minuten und 2,4 Stunden auf 1000 gemischten Seiten.

    ``use_cache`` überspringt die Extraktion, wenn dieselbe Datei mit
    denselben OCR-Einstellungen schon einmal verarbeitet wurde. ``file_hash``
    kann übergeben werden, wenn der Aufrufer ihn ohnehin gebildet hat — der
    Ingest tut das für die Änderungserkennung.
    """
    file_path = Path(path).expanduser()

    # probe() läuft vor der Cache-Prüfung, obwohl der Cache sie ersetzen
    # könnte: erst danach steht der OCR-Modus fest, und ohne ihn ist der
    # Schlüssel nicht eindeutig. Ein früherer Lauf mit --no-ocr auf einer
    # Scan-Datei würde sonst sein leeres Ergebnis für einen Lauf mit
    # automatischer Erkennung liefern. Die Voranalyse kostet ~0,05 s.
    analysis = probe(file_path)

    if ocr is None:
        ocr = bool(analysis.pages_needing_ocr)

    cache_key: str | None = None
    if use_cache:
        try:
            cache_key = file_hash or _file_hash(file_path)
        except OSError as exc:
            logger.debug("Hash für %s nicht bildbar: %s", file_path, exc)
        if cache_key:
            key = _cache_key(cache_key, ocr=ocr, ocr_langs=ocr_langs)
            if cached := _read_cache(key):
                logger.debug("Extraktion aus dem Cache: %s", file_path.name)
                return ConvertedDocument(
                    path=file_path,
                    format=cached["format"],
                    markdown=cached["markdown"],
                    ocr_used=cached["ocr_used"],
                    page_count=cached["page_count"],
                    duration_seconds=0.0,
                    warnings=list(cached.get("warnings", []))
                    + ["aus dem Extraktions-Cache"],
                )

    warnings = list(analysis.warnings)
    if analysis.sparse_pages:
        warnings.append(
            f"{len(analysis.sparse_pages)} Seite(n) ohne Text und ohne Bild — "
            "übersprungen, dort ist nichts zu holen"
        )

    started = time.time()

    suffix = file_path.suffix.lower()
    if suffix not in DOCLING_SUFFIXES:
        # Markdown und Text sind schon Text; Docling brächte hier nichts
        # außer Ladezeit. Gecacht wird trotzdem — der Gewinn ist bei diesen
        # Formaten gering, aber ein Sonderfall im Cache wäre eine Fehlerquelle
        # mehr als die paar Kilobyte wert.
        document = ConvertedDocument(
            path=file_path,
            format=analysis.format,
            markdown=analysis.text,
            ocr_used=False,
            page_count=len(analysis.pages),
            duration_seconds=time.time() - started,
            warnings=warnings,
        )
        _store_in_cache(document, cache_key, ocr=False, ocr_langs=ocr_langs)
        return document

    converter = _get_converter(ocr=ocr, ocr_langs=ocr_langs)
    try:
        result = converter.convert(str(file_path))
        markdown = result.document.export_to_markdown()
    except Exception as exc:
        raise ExtractionError(f"Docling-Konvertierung fehlgeschlagen: {exc}") from exc

    duration = time.time() - started

    # Wenn Docling weniger Text liefert als der rohe Text-Layer, ist bei der
    # Layout-Analyse etwas verloren gegangen. Sichtbar machen, nicht raten.
    if analysis.char_count > 500 and len(markdown) < analysis.char_count * 0.5:
        warnings.append(
            f"Docling lieferte {len(markdown)} Zeichen, der rohe Text-Layer "
            f"hatte {analysis.char_count} — möglicher Verlust bei der "
            "Layout-Analyse"
        )

    document = ConvertedDocument(
        path=file_path,
        format=analysis.format,
        markdown=markdown,
        ocr_used=ocr,
        page_count=len(analysis.pages),
        duration_seconds=duration,
        warnings=warnings,
    )
    _store_in_cache(document, cache_key, ocr=ocr, ocr_langs=ocr_langs)
    return document


def _store_in_cache(
    document: ConvertedDocument,
    file_hash: str | None,
    *,
    ocr: bool,
    ocr_langs: tuple[str, ...],
) -> None:
    """Ergebnis für einen späteren Lauf ablegen.

    ``file_hash`` ist None, wenn der Cache abgeschaltet war oder der Hash sich
    nicht bilden ließ — dann wird nichts geschrieben.
    """
    if not file_hash:
        return
    _write_cache(
        _cache_key(file_hash, ocr=ocr, ocr_langs=ocr_langs),
        {
            "format": document.format,
            "markdown": document.markdown,
            "ocr_used": document.ocr_used,
            "page_count": document.page_count,
            "warnings": document.warnings,
        },
    )
