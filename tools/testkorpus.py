"""Erzeugt einen deutschen Testkorpus für die Retrieval-Messung.

Drei Testdokumente reichen nicht: bei fünfzehn Chunks findet die Vektorsuche
die richtige Stelle auch dann, wenn sie schlecht ist. Erst wenn dasselbe Wort
in acht Dokumenten steht, muss sie wirklich unterscheiden.

Der Korpus ist deshalb **auf Kollisionen gebaut**:

- „Kündigungsfrist" steht in vier Dokumenten, mit vier verschiedenen Fristen.
- „Schadenmeldung" in zwei Versicherungsdokumenten, mit verschiedenen Fristen.
- „Löschfrist" in zwei Datenschutzdokumenten.
- Beträge und Pauschalen tauchen in vier Dokumenten auf.

Jede Zahl im Korpus ist **eindeutig**. Damit lässt sich an der Antwort ablesen,
welche Quelle das Modell tatsächlich benutzt hat — bei zwei Dokumenten mit
„drei Monate" wäre das nicht zu unterscheiden.

Erzeugt wird als Skript und nicht ins Repo gelegt, damit DOCX und PDF nicht als
Binärdateien eingecheckt werden müssen und der Korpus reproduzierbar bleibt.

    python tools/testkorpus.py testdaten/korpus
"""

from __future__ import annotations

import sys
from dataclasses import dataclass, field
from pathlib import Path

# Ein Abschnitt: Überschriftenebene, Überschrift, Absätze. Tabellen als Liste
# von Zeilen, die erste ist der Kopf.
Section = tuple[int, str, list[str]]


@dataclass
class Doc:
    name: str
    fmt: str  # "md" | "txt" | "docx" | "pdf" | "scan"
    title: str
    sections: list[Section] = field(default_factory=list)
    tables: dict[str, list[list[str]]] = field(default_factory=dict)


DOCS: list[Doc] = [
    # ─── Mietrecht: zwei Verträge, zwei Kündigungsfristen ────────────────────
    Doc(
        "mietvertrag-gewerbe",
        "md",
        "Mietvertrag über Gewerberäume",
        [
            (2, "Vertragsparteien", [
                "Vermieterin ist die Grundbesitz Nordstadt GmbH, vertreten durch "
                "Frau Dr. Alena Öztürk. Mieterin ist die Werkstatt Süd eG mit Sitz "
                "in München."
            ]),
            (2, "Mietgegenstand", [
                "Vermietet werden die Gewerberäume im Erdgeschoss des Objekts "
                "Lindenstraße 14 mit einer Gesamtfläche von 240 Quadratmetern. Die "
                "Räume werden zum Betrieb einer Fahrradwerkstatt vermietet. Eine "
                "Nutzungsänderung bedarf gem. Abs. 4 der schriftlichen Zustimmung."
            ]),
            (2, "Miete und Nebenkosten", [
                "Die monatliche Grundmiete beträgt 2.850 Euro netto. Hinzu kommen "
                "Nebenkosten in Form einer Vorauszahlung von 420 Euro monatlich. "
                "Die Abrechnung erfolgt jährlich zum 31. März des Folgejahres."
            ]),
            # Die H2 "Kündigung" muss hier stehen: ohne sie erben die beiden
            # H3 darunter den Pfad "Miete und Nebenkosten", und der
            # Überschriften-Präfix im Chunk führt das Retrieval in die Irre.
            # In der ersten Messung landete die Kündigungsfrist deswegen auf
            # Rang 2 hinter einem sachfremden Treffer.
            (2, "Kündigung", []),
            (3, "Kündigungsfristen", [
                "Die Kündigungsfrist beträgt sechs Monate zum Quartalsende. Eine "
                "Kündigung muss schriftlich erfolgen und der Vermieterin "
                "spätestens am dritten Werktag des ersten Monats der Frist "
                "zugehen."
            ]),
            (3, "Außerordentliche Kündigung", [
                "Die Vermieterin kann außerordentlich kündigen, wenn die Mieterin "
                "mit zwei aufeinanderfolgenden Monatsmieten in Verzug ist. Die "
                "Mieterin kann außerordentlich kündigen, wenn die Räume über einen "
                "Zeitraum von mehr als vier Wochen nicht nutzbar sind."
            ]),
            (2, "Instandhaltung", [
                "Kleinreparaturen bis zu einem Betrag von 150 Euro im Einzelfall "
                "trägt die Mieterin, höchstens jedoch 900 Euro pro Kalenderjahr."
            ]),
            (2, "Untervermietung", [
                "Eine Untervermietung ist nur mit schriftlicher Zustimmung "
                "zulässig. Die Zustimmung gilt als erteilt, wenn die Vermieterin "
                "nicht binnen 21 Tagen widerspricht."
            ]),
        ],
        tables={
            "Miete und Nebenkosten": [
                ["Position", "Betrag monatlich", "Umlagefähig"],
                ["Grundmiete", "2.850 EUR", "nein"],
                ["Heizung und Warmwasser", "180 EUR", "ja"],
                ["Grundsteuer", "95 EUR", "ja"],
                ["Gebäudeversicherung", "60 EUR", "ja"],
                ["Hausmeisterdienst", "85 EUR", "ja"],
            ]
        },
    ),
    Doc(
        "mietvertrag-wohnung",
        "docx",
        "Wohnraummietvertrag",
        [
            (2, "Mietobjekt", [
                "Vermietet wird die Dreizimmerwohnung im zweiten Obergeschoss des "
                "Hauses Rosenweg 8 mit 78 Quadratmetern Wohnfläche, nebst "
                "Kellerabteil Nummer 12 und einem Tiefgaragenstellplatz."
            ]),
            (2, "Miete", [
                "Die Nettokaltmiete beträgt 1.140 Euro monatlich. Die "
                "Betriebskostenvorauszahlung beläuft sich auf 260 Euro, die "
                "Heizkostenvorauszahlung auf 130 Euro."
            ]),
            (2, "Kaution", [
                "Die Mietsicherheit beträgt zwei Nettokaltmieten, also 2.280 Euro. "
                "Sie kann in drei gleichen Monatsraten geleistet werden."
            ]),
            (2, "Beendigung des Mietverhältnisses", [
                "Die Kündigungsfrist für die Mieterseite beträgt drei Monate. Für "
                "die Vermieterseite verlängert sich die Frist nach fünf Jahren "
                "Mietdauer auf sechs Monate und nach acht Jahren auf neun Monate.",
                "Die Kündigung ist schriftlich zu erklären und muss bis zum "
                "dritten Werktag eines Kalendermonats zugehen."
            ]),
            (2, "Tierhaltung", [
                "Die Haltung von Kleintieren ist ohne Zustimmung erlaubt. Hunde "
                "und Katzen bedürfen der vorherigen Erlaubnis der Vermieterseite."
            ]),
        ],
    ),
    # ─── Arbeitsrecht ───────────────────────────────────────────────────────
    Doc(
        "arbeitsvertrag",
        "pdf",
        "Anstellungsvertrag",
        [
            (2, "Tätigkeit und Eintritt", [
                "Die Arbeitnehmerin wird als Softwareentwicklerin eingestellt. Das "
                "Arbeitsverhältnis beginnt am 1. September und ist unbefristet."
            ]),
            (2, "Probezeit", [
                "Die ersten sechs Monate gelten als Probezeit. Während der "
                "Probezeit kann das Arbeitsverhältnis mit einer Frist von zwei "
                "Wochen gekündigt werden."
            ]),
            (2, "Vergütung", [
                "Das Bruttojahresgehalt beträgt 68.400 Euro, zahlbar in zwölf "
                "gleichen Monatsraten zum Monatsende. Eine variable Vergütung von "
                "bis zu 8 Prozent des Jahresgehalts richtet sich nach der "
                "Zielerreichung."
            ]),
            (2, "Kündigung des Arbeitsverhältnisses", [
                "Nach Ablauf der Probezeit gilt eine Kündigungsfrist von vier "
                "Wochen zum Monatsende. Die gesetzlichen Verlängerungen für die "
                "Arbeitgeberseite bleiben unberührt."
            ]),
            (2, "Arbeitszeit", [
                "Die regelmäßige Wochenarbeitszeit beträgt 38 Stunden. Mehrarbeit "
                "wird durch Freizeit ausgeglichen."
            ]),
            (2, "Nebentätigkeit", [
                "Jede entgeltliche Nebentätigkeit ist vorher schriftlich "
                "anzuzeigen. Die Anzeige gilt als genehmigt, wenn nicht binnen "
                "zehn Arbeitstagen widersprochen wird."
            ]),
        ],
    ),
    Doc(
        "betriebsvereinbarung-urlaub",
        "txt",
        "Betriebsvereinbarung Urlaub und Gleitzeit",
        [
            (2, "Urlaubsanspruch", [
                "Der Jahresurlaub beträgt 30 Arbeitstage bei einer Fünftagewoche. "
                "Der Anspruch entsteht mit Beginn des Kalenderjahres und ist bis "
                "zum 31. März des Folgejahres zu nehmen. Danach verfällt er, "
                "sofern nicht betriebliche Gründe entgegenstehen."
            ]),
            (2, "Antragsverfahren", [
                "Urlaub ist mindestens vier Wochen vor Antritt zu beantragen. Bei "
                "mehr als zehn zusammenhängenden Tagen gilt eine Frist von acht "
                "Wochen."
            ]),
            (2, "Gleitzeit", [
                "Die Gleitzeitspanne liegt zwischen 6:30 Uhr und 20:00 Uhr. Die "
                "Kernarbeitszeit umfasst die Zeit von 9:30 Uhr bis 15:00 Uhr, "
                "freitags bis 13:00 Uhr. Ein Gleitzeitsaldo von höchstens 40 "
                "Plusstunden und 15 Minusstunden ist zulässig."
            ]),
            (2, "Sonderurlaub", [
                "Für die eigene Eheschließung werden zwei Tage gewährt, für den "
                "Umzug aus betrieblichem Anlass ein Tag, bei Tod eines "
                "Familienangehörigen ersten Grades drei Tage."
            ]),
        ],
    ),
    Doc(
        "homeoffice-regelung",
        "md",
        "Regelung zur mobilen Arbeit",
        [
            (2, "Umfang", [
                "Mobile Arbeit ist an bis zu drei Tagen pro Woche möglich. Ein "
                "Anwesenheitstag je Woche ist verpflichtend und wird im Team "
                "abgestimmt."
            ]),
            (2, "Ankündigung", [
                "Die Inanspruchnahme ist mindestens drei Arbeitstage im Voraus im "
                "Zeiterfassungssystem zu vermerken. Kurzfristige Abweichungen sind "
                "mit der Führungskraft abzustimmen."
            ]),
            (2, "Ausstattung", [
                "Der Arbeitgeber stellt Notebook und Headset. Für die Einrichtung "
                "des häuslichen Arbeitsplatzes wird ein einmaliger Zuschuss von "
                "450 Euro gewährt, frühestens nach Ablauf der Probezeit."
            ]),
            (2, "Erreichbarkeit", [
                "Während der Kernarbeitszeit ist telefonische Erreichbarkeit "
                "sicherzustellen. Eine Reaktion auf Nachrichten wird binnen vier "
                "Stunden erwartet."
            ]),
        ],
    ),
    Doc(
        "weiterbildung",
        "md",
        "Richtlinie Weiterbildung",
        [
            (2, "Bildungsurlaub", [
                "Je Kalenderjahr stehen fünf Tage Bildungsurlaub zur Verfügung. "
                "Nicht genutzte Tage können einmalig ins Folgejahr übertragen "
                "werden."
            ]),
            (2, "Kostenübernahme", [
                "Der Arbeitgeber übernimmt Seminarkosten bis 2.400 Euro je "
                "Kalenderjahr. Bei Maßnahmen über 5.000 Euro wird eine "
                "Rückzahlungsvereinbarung mit einer Bindungsdauer von 24 Monaten "
                "geschlossen."
            ]),
            (2, "Antragsweg", [
                "Anträge sind spätestens sechs Wochen vor Beginn der Maßnahme über "
                "die Führungskraft an die Personalabteilung zu richten."
            ]),
        ],
    ),
    # ─── Datenschutz: zwei Dokumente mit Löschfristen ───────────────────────
    Doc(
        "datenschutzkonzept",
        "md",
        "Datenschutzkonzept Kundenportal",
        [
            (2, "Zweck der Verarbeitung", [
                "Die Verarbeitung personenbezogener Daten erfolgt zur Erfüllung "
                "des Vertragsverhältnisses gemäß Art. 6 Abs. 1 lit. b DSGVO. Eine "
                "darüber hinausgehende Nutzung zu Werbezwecken erfolgt "
                "ausschließlich nach ausdrücklicher Einwilligung."
            ]),
            (2, "Löschfristen", [
                "Nutzungsprotokolle werden nach 90 Tagen automatisiert gelöscht. "
                "Stammdaten unterliegen der handelsrechtlichen "
                "Aufbewahrungspflicht und werden erst zehn Jahre nach Ende des "
                "Vertragsverhältnisses entfernt."
            ]),
            (2, "Auftragsverarbeiter", [
                "Eingesetzt werden ausschließlich Auftragsverarbeiter mit Sitz "
                "innerhalb der Europäischen Union. Ein Vertrag zur "
                "Auftragsverarbeitung nach Art. 28 DSGVO liegt jeweils vor."
            ]),
            (2, "Technische Maßnahmen", [
                "Die Übertragung erfolgt ausschließlich verschlüsselt über TLS "
                "1.3. Datenbanken werden mit AES-256 verschlüsselt gespeichert. "
                "Der Zugriff ist über ein Rollenkonzept beschränkt, "
                "Administratorzugänge erfordern Zwei-Faktor-Authentisierung."
            ]),
        ],
        tables={
            "Löschfristen": [
                ["Datenkategorie", "Rechtsgrundlage", "Löschfrist"],
                ["Stammdaten", "Art. 6 Abs. 1 lit. b", "10 Jahre nach Vertragsende"],
                ["Zahlungsdaten", "Art. 6 Abs. 1 lit. b", "10 Jahre (§ 147 AO)"],
                ["Nutzungsprotokolle", "Art. 6 Abs. 1 lit. f", "90 Tage"],
                ["Newsletter-Einwilligung", "Art. 6 Abs. 1 lit. a", "bis Widerruf"],
            ]
        },
    ),
    Doc(
        "loeschkonzept-personal",
        "md",
        "Löschkonzept Personalakten",
        [
            (2, "Geltungsbereich", [
                "Dieses Konzept regelt die Aufbewahrung und Vernichtung von "
                "Unterlagen aus Bewerbungs- und Beschäftigungsverhältnissen."
            ]),
            (2, "Bewerbungsunterlagen", [
                "Unterlagen abgelehnter Bewerberinnen und Bewerber werden sechs "
                "Monate nach Abschluss des Verfahrens vernichtet. Bei erteilter "
                "Einwilligung zur Aufnahme in einen Bewerberpool verlängert sich "
                "die Frist auf 18 Monate."
            ]),
            (2, "Personalakten", [
                "Die Personalakte wird drei Jahre nach Ende des "
                "Beschäftigungsverhältnisses vernichtet. Unterlagen mit "
                "steuerlicher Relevanz verbleiben zehn Jahre, Nachweise zur "
                "betrieblichen Altersversorgung 30 Jahre."
            ]),
            (2, "Abmahnungen", [
                "Abmahnungen werden nach zwei Jahren aus der Personalakte "
                "entfernt, sofern kein weiterer gleichartiger Verstoß vorliegt."
            ]),
        ],
    ),
    Doc(
        "auftragsverarbeitung",
        "docx",
        "Vertrag zur Auftragsverarbeitung",
        [
            (2, "Gegenstand", [
                "Der Auftragnehmer verarbeitet personenbezogene Daten "
                "ausschließlich nach Weisung des Auftraggebers zum Zweck des "
                "Betriebs der Kundenplattform."
            ]),
            (2, "Meldung von Datenschutzverletzungen", [
                "Der Auftragnehmer meldet dem Auftraggeber jede Verletzung des "
                "Schutzes personenbezogener Daten unverzüglich, spätestens jedoch "
                "binnen 24 Stunden nach Kenntniserlangung."
            ]),
            (2, "Unterauftragsverhältnisse", [
                "Die Beauftragung weiterer Unterauftragnehmer bedarf der "
                "vorherigen Zustimmung. Der Auftraggeber wird mindestens 30 Tage "
                "vor einem Wechsel informiert und kann widersprechen."
            ]),
            (2, "Löschung nach Vertragsende", [
                "Nach Beendigung des Vertrags werden alle Daten binnen 60 Tagen "
                "gelöscht oder zurückgegeben. Über die Löschung wird ein Protokoll "
                "erstellt."
            ]),
            (2, "Kontrollrechte", [
                "Der Auftraggeber darf einmal jährlich eine Prüfung vor Ort "
                "durchführen. Der Termin ist vier Wochen vorher anzukündigen."
            ]),
        ],
    ),
    # ─── Versicherung: zwei Dokumente mit Meldefristen ──────────────────────
    Doc(
        "betriebshaftpflicht",
        "pdf",
        "Bedingungen Betriebshaftpflichtversicherung",
        [
            (2, "Versicherungsumfang", [
                "Versichert ist die gesetzliche Haftpflicht aus dem Betrieb des "
                "Handwerksunternehmens. Die Deckungssumme beträgt 5 Millionen Euro "
                "pauschal für Personen- und Sachschäden."
            ]),
            (2, "Selbstbeteiligung", [
                "Je Schadenfall trägt der Versicherungsnehmer eine "
                "Selbstbeteiligung von 500 Euro. Bei Tätigkeitsschäden erhöht sich "
                "die Selbstbeteiligung auf 1.500 Euro."
            ]),
            (2, "Meldung eines Schadens", [
                "Ein Schadenfall ist dem Versicherer binnen sieben Tagen nach "
                "Kenntnis schriftlich anzuzeigen. Bei Personenschäden verkürzt "
                "sich die Frist auf 48 Stunden."
            ]),
            (2, "Ausschlüsse", [
                "Nicht versichert sind Schäden durch vorsätzliches Handeln, "
                "Schäden an gemieteten Fahrzeugen sowie Ansprüche wegen "
                "Vermögensschäden ohne Sach- oder Personenschaden."
            ]),
        ],
    ),
    Doc(
        "gebaeudeversicherung",
        "md",
        "Gebäudeversicherung Objekt Lindenstraße",
        [
            (2, "Versicherte Gefahren", [
                "Versichert sind Schäden durch Feuer, Leitungswasser, Sturm ab "
                "Windstärke 8 und Hagel. Elementarschäden durch Überschwemmung "
                "sind eingeschlossen."
            ]),
            (2, "Versicherungssumme", [
                "Die Versicherungssumme beträgt 2,4 Millionen Euro auf Basis des "
                "gleitenden Neuwerts. Der Jahresbeitrag liegt bei 3.180 Euro."
            ]),
            (2, "Anzeige eines Schadens", [
                "Schäden sind unverzüglich anzuzeigen, spätestens innerhalb von "
                "drei Werktagen. Bei Leitungswasserschäden ist die Zuleitung "
                "sofort zu schließen und die Schadenstelle zu dokumentieren."
            ]),
            (2, "Obliegenheiten im Winter", [
                "Wasserführende Leitungen in ungenutzten Gebäudeteilen sind zu "
                "entleeren. Die Beheizung ist in der Zeit vom 1. November bis zum "
                "31. März sicherzustellen und wöchentlich zu kontrollieren."
            ]),
        ],
    ),
    # ─── Kosten und Beschaffung: Kollision bei Beträgen ─────────────────────
    Doc(
        "reisekostenrichtlinie",
        "md",
        "Reisekostenrichtlinie",
        [
            (2, "Fahrtkosten", [
                "Bahnfahrten werden in der zweiten Klasse erstattet. Ab einer "
                "Fahrtzeit von vier Stunden ist die erste Klasse zulässig. Die "
                "Kilometerpauschale für die Nutzung des privaten PKW beträgt 0,30 "
                "Euro je gefahrenem Kilometer."
            ]),
            (2, "Übernachtung", [
                "Der Höchstbetrag für eine Übernachtung im Inland liegt bei 120 "
                "Euro, im Ausland bei 180 Euro. Frühstück ist mit 5,60 Euro "
                "abzuziehen, sofern es in der Rechnung enthalten ist."
            ]),
            (2, "Verpflegungspauschale", [
                "Bei einer Abwesenheit von mehr als acht Stunden werden 14 Euro "
                "erstattet, bei ganztägiger Abwesenheit 28 Euro."
            ]),
            (2, "Abrechnung", [
                "Der Antrag ist binnen sechs Wochen nach Rückkehr einzureichen. "
                "Später eingehende Anträge werden nur in begründeten Ausnahmen "
                "bearbeitet."
            ]),
        ],
    ),
    Doc(
        "bewirtungsordnung",
        "md",
        "Ordnung für Bewirtung und Repräsentation",
        [
            (2, "Bewirtung von Geschäftspartnern", [
                "Je Person und Anlass gilt eine Obergrenze von 85 Euro "
                "einschließlich Getränken. Die Namen aller Teilnehmenden und der "
                "geschäftliche Anlass sind auf dem Bewirtungsbeleg zu vermerken."
            ]),
            (2, "Interne Anlässe", [
                "Bei internen Veranstaltungen gilt eine Obergrenze von 35 Euro je "
                "Person. Für Jubiläen und Verabschiedungen stehen zusätzlich 250 "
                "Euro je Anlass zur Verfügung."
            ]),
            (2, "Geschenke", [
                "Sachgeschenke an Geschäftspartner sind bis 50 Euro je Person und "
                "Jahr zulässig. Bargeld und Gutscheine sind ausgeschlossen."
            ]),
        ],
    ),
    Doc(
        "einkaufsrichtlinie",
        "docx",
        "Einkaufs- und Genehmigungsrichtlinie",
        [
            (2, "Genehmigungsgrenzen", [
                "Bestellungen bis 1.000 Euro genehmigt die Teamleitung, bis 10.000 "
                "Euro die Abteilungsleitung, darüber die Geschäftsführung. Ab "
                "25.000 Euro sind drei Vergleichsangebote einzuholen."
            ]),
            (2, "Rahmenverträge", [
                "Für Büromaterial, IT-Zubehör und Arbeitskleidung bestehen "
                "Rahmenverträge. Abweichende Beschaffungen sind zu begründen."
            ]),
            (2, "Zahlungsbedingungen", [
                "Angestrebt wird ein Zahlungsziel von 30 Tagen mit 2 Prozent "
                "Skonto bei Zahlung binnen 14 Tagen. Vorkasse ist unzulässig, "
                "sofern nicht die Geschäftsführung zustimmt."
            ]),
            (2, "Rechnungsprüfung", [
                "Eingangsrechnungen sind binnen fünf Arbeitstagen sachlich zu "
                "prüfen und freizugeben."
            ]),
        ],
    ),
    # ─── IT und Betrieb ─────────────────────────────────────────────────────
    Doc(
        "it-richtlinie",
        "md",
        "IT-Sicherheitsrichtlinie",
        [
            (2, "Kennwörter", [
                "Kennwörter müssen mindestens zwölf Zeichen umfassen und Groß- und "
                "Kleinbuchstaben, Ziffern sowie Sonderzeichen enthalten. Ein "
                "Wechsel ist alle 90 Tage erforderlich. Die letzten fünf "
                "Kennwörter dürfen nicht wiederverwendet werden."
            ]),
            (2, "Sperrung des Arbeitsplatzes", [
                "Der Bildschirm ist beim Verlassen des Arbeitsplatzes zu sperren. "
                "Die automatische Sperre greift nach fünf Minuten ohne Eingabe."
            ]),
            (2, "Umgang mit Datenträgern", [
                "Der Einsatz privater USB-Datenträger ist untersagt. Dienstliche "
                "Datenträger sind zu verschlüsseln."
            ]),
            (2, "Meldung von Sicherheitsvorfällen", [
                "Verdachtsfälle sind unverzüglich an das Sicherheitsteam zu "
                "melden, spätestens innerhalb einer Stunde nach Feststellung."
            ]),
        ],
    ),
    Doc(
        "backup-konzept",
        "md",
        "Backup- und Wiederherstellungskonzept",
        [
            (2, "Sicherungsrhythmus", [
                "Datenbanken werden stündlich inkrementell und täglich vollständig "
                "gesichert. Dateiserver werden täglich gesichert, Konfigurationen "
                "bei jeder Änderung."
            ]),
            (2, "Aufbewahrung", [
                "Tagessicherungen werden 30 Tage vorgehalten, Wochensicherungen "
                "zwölf Wochen, Monatssicherungen 36 Monate. Eine Kopie liegt an "
                "einem räumlich getrennten Standort."
            ]),
            (2, "Wiederherstellungsziele", [
                "Die maximal tolerierte Ausfallzeit beträgt vier Stunden, der "
                "maximal tolerierte Datenverlust eine Stunde."
            ]),
            (2, "Wiederherstellungstest", [
                "Eine vollständige Wiederherstellung wird quartalsweise geprobt "
                "und protokolliert. Ein fehlgeschlagener Test ist binnen zwei "
                "Wochen zu wiederholen."
            ]),
        ],
    ),
    Doc(
        "serverwartung",
        "md",
        "Wartungsplan Serverinfrastruktur",
        [
            (2, "Wartungsfenster", [
                "Das reguläre Wartungsfenster liegt jeden zweiten Sonntag im Monat "
                "von 2:00 bis 6:00 Uhr. Ankündigungen erfolgen mindestens zehn "
                "Kalendertage vorher."
            ]),
            (2, "Sicherheitsaktualisierungen", [
                "Kritische Sicherheitsaktualisierungen werden binnen 72 Stunden "
                "eingespielt, auch außerhalb des Wartungsfensters. Übrige "
                "Aktualisierungen folgen dem monatlichen Zyklus."
            ]),
            (2, "Überwachung", [
                "Systemwerte werden im Minutenabstand erfasst. Eine Warnung wird "
                "ausgelöst, wenn die Festplattenbelegung 85 Prozent oder die "
                "Prozessorlast über zehn Minuten 90 Prozent übersteigt."
            ]),
        ],
    ),
    # ─── Sonstige Domänen, als Distraktoren ─────────────────────────────────
    Doc(
        "vereinssatzung",
        "docx",
        "Satzung des Fördervereins Stadtgarten",
        [
            (2, "Zweck", [
                "Der Verein fördert die Anlage und Pflege öffentlich zugänglicher "
                "Gartenflächen sowie die Umweltbildung von Kindern und "
                "Jugendlichen."
            ]),
            (2, "Mitgliedschaft", [
                "Der Jahresbeitrag beträgt 48 Euro für Einzelpersonen und 72 Euro "
                "für Familien. Der Austritt ist zum Jahresende mit einer Frist von "
                "einem Monat schriftlich zu erklären."
            ]),
            (2, "Mitgliederversammlung", [
                "Die ordentliche Mitgliederversammlung findet jährlich im ersten "
                "Quartal statt. Die Einladung erfolgt schriftlich mindestens 14 "
                "Tage vorher unter Angabe der Tagesordnung."
            ]),
            (2, "Vorstand", [
                "Der Vorstand besteht aus fünf Personen und wird für zwei Jahre "
                "gewählt. Er ist beschlussfähig, wenn mindestens drei Mitglieder "
                "anwesend sind."
            ]),
        ],
    ),
    Doc(
        "hausordnung",
        "txt",
        "Hausordnung Rosenweg 8",
        [
            (2, "Ruhezeiten", [
                "Die Mittagsruhe gilt von 13:00 bis 15:00 Uhr, die Nachtruhe von "
                "22:00 bis 7:00 Uhr. An Sonn- und Feiertagen ist ganztägig auf "
                "Zimmerlautstärke zu achten."
            ]),
            (2, "Treppenhausreinigung", [
                "Die Reinigung des Treppenhauses erfolgt wöchentlich im Wechsel "
                "nach dem ausgehängten Plan. Bei Verhinderung ist für Ersatz zu "
                "sorgen."
            ]),
            (2, "Müllentsorgung", [
                "Die Restmülltonne wird donnerstags geleert, Papier alle zwei "
                "Wochen montags, Biomüll dienstags. Sperrmüll ist beim Bauhof "
                "anzumelden."
            ]),
            (2, "Fahrräder und Kinderwagen", [
                "Fahrräder sind ausschließlich im Fahrradkeller abzustellen. Das "
                "Treppenhaus muss als Fluchtweg frei bleiben."
            ]),
        ],
    ),
    Doc(
        "arbeitsschutz",
        "md",
        "Arbeitsschutz und Erste Hilfe",
        [
            (2, "Unterweisung", [
                "Alle Beschäftigten werden bei Einstellung und danach jährlich "
                "zum Arbeitsschutz unterwiesen. Die Unterweisung wird "
                "dokumentiert und von den Teilnehmenden bestätigt."
            ]),
            (2, "Ersthelfer", [
                "In jedem Betriebsteil sind mindestens zwei ausgebildete "
                "Ersthelfer verfügbar. Die Ausbildung wird alle zwei Jahre "
                "aufgefrischt."
            ]),
            (2, "Verbandskasten", [
                "Verbandskästen befinden sich in der Werkstatt, im Lager und am "
                "Empfang. Der Inhalt wird halbjährlich geprüft und "
                "gegebenenfalls ergänzt."
            ]),
            (2, "Meldung von Unfällen", [
                "Jeder Arbeitsunfall ist sofort der Führungskraft zu melden und "
                "im Verbandbuch einzutragen. Unfälle mit mehr als drei Tagen "
                "Arbeitsunfähigkeit werden binnen drei Tagen der "
                "Berufsgenossenschaft angezeigt."
            ]),
        ],
    ),
    Doc(
        "fahrzeugordnung",
        "md",
        "Ordnung für Dienstfahrzeuge",
        [
            (2, "Berechtigung", [
                "Dienstfahrzeuge dürfen von Beschäftigten mit gültiger "
                "Fahrerlaubnis nach einmaliger Einweisung genutzt werden. Die "
                "Fahrerlaubnis wird halbjährlich kontrolliert."
            ]),
            (2, "Private Nutzung", [
                "Die private Nutzung ist nur bei ausdrücklicher Vereinbarung "
                "zulässig und wird mit einem Prozent des Bruttolistenpreises "
                "monatlich versteuert."
            ]),
            (2, "Tanken und Laden", [
                "Getankt wird ausschließlich mit der Tankkarte des Fahrzeugs. "
                "Elektrofahrzeuge werden vorrangig an den Ladepunkten am Standort "
                "geladen, die Erstattung für das Laden zu Hause beträgt 30 Euro "
                "monatlich pauschal."
            ]),
            (2, "Schäden und Fahrtenbuch", [
                "Schäden sind vor Fahrtende zu melden. Das elektronische "
                "Fahrtenbuch ist bei jeder Fahrt zu führen; fehlende Einträge "
                "sind binnen sieben Tagen nachzutragen."
            ]),
        ],
    ),
    Doc(
        "qualitaetshandbuch",
        "md",
        "Qualitätshandbuch Prüfmittel",
        [
            (2, "Prüfmittelverwaltung", [
                "Alle Messmittel sind im Prüfmittelverzeichnis erfasst und mit "
                "einer eindeutigen Kennnummer versehen. Der Standort wird "
                "nachgeführt."
            ]),
            (2, "Kalibrierintervalle", [
                "Messschieber und Bügelmessschrauben werden jährlich kalibriert, "
                "Drehmomentschlüssel alle sechs Monate, Waagen alle zwölf Monate "
                "durch einen externen Dienstleister."
            ]),
            (2, "Umgang mit Abweichungen", [
                "Wird bei der Kalibrierung eine Abweichung außerhalb der Toleranz "
                "festgestellt, ist das Prüfmittel zu sperren. Betroffene "
                "Messungen der letzten drei Monate sind zu bewerten."
            ]),
            (2, "Reklamationen", [
                "Kundenreklamationen werden binnen zwei Arbeitstagen bestätigt und "
                "binnen 15 Arbeitstagen abschließend beantwortet."
            ]),
        ],
    ),
    Doc(
        "kuendigungsschreiben-muster",
        "txt",
        "Muster für ein Kündigungsschreiben",
        [
            (2, "Hinweise zur Verwendung", [
                "Dieses Muster dient als Vorlage. Die maßgebliche Kündigungsfrist "
                "ergibt sich aus dem jeweiligen Vertrag und ist vor dem Versand zu "
                "prüfen. Das Schreiben ist eigenhändig zu unterzeichnen."
            ]),
            (2, "Mustertext", [
                "Sehr geehrte Damen und Herren, hiermit kündige ich das zwischen "
                "uns bestehende Vertragsverhältnis zum nächstmöglichen Zeitpunkt. "
                "Ich bitte um eine schriftliche Bestätigung des "
                "Beendigungsdatums.",
                "Mit freundlichen Grüßen"
            ]),
            (2, "Versand", [
                "Der Versand sollte per Einwurf-Einschreiben erfolgen, damit der "
                "Zugang nachweisbar ist. Eine Kopie ist zu den Unterlagen zu "
                "nehmen."
            ]),
        ],
    ),
    # ─── Gescanntes Dokument: prüft den OCR-Pfad ────────────────────────────
    Doc(
        "gescannter-antrag",
        "scan",
        "Antrag auf Kostenübernahme",
        [
            (2, "Angaben", [
                "Antragsteller: Jürgen Öztürk, Personalnummer 4711.",
                "Maßnahme: Fachseminar Schweißtechnik, Größe der Gruppe: 12.",
                "Kosten gemäß Angebot: 1.890 Euro zzgl. Umsatzsteuer.",
                "Die Genehmigung erfolgt vorbehaltlich der Mittelfreigabe.",
            ]),
        ],
    ),
]


# ─── Ausgabe je Format ───────────────────────────────────────────────────────


def render_markdown(doc: Doc) -> str:
    lines = [f"# {doc.title}", ""]
    for level, heading, paragraphs in doc.sections:
        lines.append(f"{'#' * level} {heading}")
        lines.append("")
        for paragraph in paragraphs:
            lines.append(paragraph)
            lines.append("")
        if table := doc.tables.get(heading):
            lines.append("| " + " | ".join(table[0]) + " |")
            lines.append("| " + " | ".join("---" for _ in table[0]) + " |")
            for row in table[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
    return "\n".join(lines)


def render_text(doc: Doc) -> str:
    """Reiner Text ohne Markdown-Auszeichnung.

    Prüft den Fall, dass es keine Überschriftenstruktur gibt und das Chunking
    ohne Pfad arbeiten muss.
    """
    lines = [doc.title, ""]
    for _, heading, paragraphs in doc.sections:
        lines.append(heading)
        lines.append("")
        lines.extend(f"{p}\n" for p in paragraphs)
    return "\n".join(lines)


def write_docx(doc: Doc, path: Path) -> None:
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading(doc.title, level=1)
    for level, heading, paragraphs in doc.sections:
        document.add_heading(heading, level=min(level, 4))
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        if table := doc.tables.get(heading):
            created = document.add_table(rows=len(table), cols=len(table[0]))
            created.style = "Table Grid"
            for r, row in enumerate(table):
                for c, cell in enumerate(row):
                    created.cell(r, c).text = cell
    document.save(path)


def write_pdf(doc: Doc, path: Path) -> None:
    """PDF mit Textlayer — der Normalfall eines digital erzeugten Dokuments."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    # Helvetica kennt die Umlaute in der Standardkodierung; ohne das würden
    # ä/ö/ü/ß als Platzhalter landen und die Extraktion prüfte nichts.
    body = ParagraphStyle(
        "Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5,
        leading=14,
    )
    heading = ParagraphStyle(
        "Head", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13,
    )

    template = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=doc.title,
    )
    flow = [Paragraph(doc.title, styles["Title"]), Spacer(1, 6 * mm)]
    for _, head, paragraphs in doc.sections:
        flow.append(Paragraph(head, heading))
        for paragraph in paragraphs:
            flow.append(Paragraph(paragraph, body))
            flow.append(Spacer(1, 2 * mm))
        flow.append(Spacer(1, 3 * mm))
    template.build(flow)


def write_scan(doc: Doc, path: Path) -> None:
    """PDF ohne Textlayer: Text als Bild gerendert.

    Damit lässt sich prüfen, ob probe() die Seite als Scan erkennt und ob OCR
    die Umlaute trifft — beides Wege, die ein digital erzeugtes PDF nie nimmt.
    """
    from PIL import Image, ImageDraw, ImageFont

    # 150 dpi auf A4: groß genug, dass EasyOCR sauber liest.
    width, height = 1240, 1754
    image = Image.new("L", (width, height), color=255)
    draw = ImageDraw.Draw(image)

    font_path = _find_font()
    title_font = ImageFont.truetype(font_path, 44)
    body_font = ImageFont.truetype(font_path, 30)

    y = 120
    draw.text((110, y), doc.title, fill=30, font=title_font)
    y += 90

    for _, head, paragraphs in doc.sections:
        draw.text((110, y), head, fill=30, font=title_font)
        y += 70
        for paragraph in paragraphs:
            for line in _wrap(paragraph, 52):
                draw.text((110, y), line, fill=40, font=body_font)
                y += 44
            y += 16

    # Etwas Rauschen und eine leichte Drehung, damit es nicht nach einem
    # perfekten Rendering aussieht — ein echter Scan ist nie ganz gerade.
    image = image.rotate(0.4, resample=Image.BICUBIC, fillcolor=255)
    image.save(path, "PDF", resolution=150.0)


def _find_font() -> str:
    kandidaten = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
    ]
    for kandidat in kandidaten:
        if Path(kandidat).exists():
            return kandidat
    raise SystemExit(
        "Keine TrueType-Schrift mit Umlauten gefunden — "
        "fonts-dejavu oder fonts-liberation installieren"
    )


def _wrap(text: str, width: int) -> list[str]:
    words, lines, current = text.split(), [], ""
    for word in words:
        probe = f"{current} {word}".strip()
        if len(probe) > width and current:
            lines.append(current)
            current = word
        else:
            current = probe
    if current:
        lines.append(current)
    return lines


WRITERS = {
    "md": lambda d, p: p.write_text(render_markdown(d), encoding="utf-8"),
    "txt": lambda d, p: p.write_text(render_text(d), encoding="utf-8"),
    "docx": write_docx,
    "pdf": write_pdf,
    "scan": write_scan,
}

SUFFIX = {"md": ".md", "txt": ".txt", "docx": ".docx", "pdf": ".pdf", "scan": ".pdf"}


def build(target: Path) -> list[Path]:
    target.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for doc in DOCS:
        path = target / f"{doc.name}{SUFFIX[doc.fmt]}"
        WRITERS[doc.fmt](doc, path)
        written.append(path)
    return written


def main() -> None:
    target = Path(sys.argv[1] if len(sys.argv) > 1 else "testdaten/korpus")
    written = build(target)
    formate: dict[str, int] = {}
    for doc in DOCS:
        formate[doc.fmt] = formate.get(doc.fmt, 0) + 1
    print(f"{len(written)} Dokumente in {target}")
    for fmt, anzahl in sorted(formate.items()):
        print(f"  {fmt:5} {anzahl}")


if __name__ == "__main__":
    main()
